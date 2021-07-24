'''
Custom Invitations as Word Documents
Say you have a text file of guest names. This guests.txt file has one name 
per line, as follows:


Prof. Plum
Miss Scarlet
Col. Mustard
Al Sweigart
Robocop
Write a program that would generate a Word document with custom invitations 
that look like Figure 13-11.

Since Python-Docx can use only those styles that already exist in the Word 
document, you will have to first add these styles to a blank Word file and 
then open that file with Python-Docx. There should be one invitation per 
page in the resulting Word document, so call add_break() to add a page 
break after the last paragraph of each invitation. This way, you will need 
to open only one Word document to print all of the invitations at once.
'''
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def text_style(paragraph, font, size, is_bold = False):
	paragraph.runs[0].font.name = font
	paragraph.runs[0].font.size = Pt(size)
	paragraph.runs[0].bold = is_bold
	paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

guests = open('guests.txt').read().split('\n')

doc = docx.Document()
aux = 0

for name in guests:
	doc.add_paragraph('It would be a pleasure to have the company of')
	text_style(doc.paragraphs[aux], 'Script MT Bold', 22)
	doc.add_paragraph(name)
	text_style(doc.paragraphs[aux + 1], 'Algerian', 20, True)
	doc.add_paragraph('At 11010 Memory Lane on the Evening of')
	text_style(doc.paragraphs[aux + 2], 'Script MT Bold', 22)
	doc.add_paragraph('April 1st')
	text_style(doc.paragraphs[aux + 3], 'Algerian', 20)
	doc.add_paragraph('at 7 o’clock')
	text_style(doc.paragraphs[aux + 4], 'Script MT Bold', 22)
	doc.add_page_break()
	aux = aux + 6

doc.save('other.docx')





